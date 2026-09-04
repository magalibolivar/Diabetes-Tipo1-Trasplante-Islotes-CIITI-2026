# -*- coding: utf-8 -*-
"""04 - Genera el paper (.docx) del estudio de labilidad glucémica / candidatura a
trasplante de islotes, en el formato de Referencia.docx (UAI/CAETI)."""
from pathlib import Path
import pandas as pd
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH as AL
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parent.parent
FIG, TAB = ROOT / "figures", ROOT / "tables"
OUT = ROOT / "paper" / "Labilidad_Glucemica_Trasplante_Islotes_DM1_UAI.docx"
OUT.parent.mkdir(parents=True, exist_ok=True)

doc = Document()
s = doc.sections[0]
s.page_width, s.page_height = Cm(21.0), Cm(29.7)
s.left_margin = s.right_margin = s.top_margin = s.bottom_margin = Cm(2.54)
normal = doc.styles["Normal"]; normal.font.name = "Times New Roman"; normal.font.size = Pt(11)
normal._element.rPr.rFonts.set(qn('w:eastAsia'), "Times New Roman")
pf = normal.paragraph_format; pf.alignment = AL.JUSTIFY; pf.space_after = Pt(6); pf.line_spacing = 1.0

def _font(run, name="Times New Roman", size=12, bold=False, italic=False):
    run.font.name = name; run.font.size = Pt(size); run.bold = bold; run.italic = italic
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)

def para(text="", align=AL.JUSTIFY, bold=False, italic=False, size=11, space_after=6, first_indent=None):
    p = doc.add_paragraph(); p.alignment = align; p.paragraph_format.space_after = Pt(space_after)
    if first_indent: p.paragraph_format.first_line_indent = Cm(first_indent)
    if text: _font(p.add_run(text), "Times New Roman", size, bold, italic)
    return p

def runs_para(segments, align=AL.JUSTIFY, space_after=6):
    p = doc.add_paragraph(); p.alignment = align; p.paragraph_format.space_after = Pt(space_after)
    for t, b, i in segments: _font(p.add_run(t), "Times New Roman", 11, b, i)
    return p

def h1(text):
    p = doc.add_paragraph(); p.alignment = AL.LEFT
    p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(6)
    _font(p.add_run(text), bold=True); return p

def h2(text):
    p = doc.add_paragraph(); p.alignment = AL.LEFT
    p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(4)
    _font(p.add_run(text), bold=True); return p

def figura(path, caption, width=15.0):
    p = doc.add_paragraph(); p.alignment = AL.CENTER; p.paragraph_format.space_before = Pt(6)
    p.add_run().add_picture(str(path), width=Cm(width))
    c = doc.add_paragraph(); c.alignment = AL.CENTER; c.paragraph_format.space_after = Pt(8)
    _font(c.add_run(caption), "Times New Roman", 11)

def tabla_csv(csv_path, caption):
    df = pd.read_csv(csv_path)
    c = doc.add_paragraph(); c.alignment = AL.CENTER
    c.paragraph_format.space_before = Pt(6); c.paragraph_format.space_after = Pt(2)
    _font(c.add_run(caption), "Times New Roman", 11, bold=True)
    cols = list(df.columns)
    t = doc.add_table(rows=1, cols=len(cols)); t.alignment = WD_TABLE_ALIGNMENT.CENTER; t.style = "Table Grid"
    for j, cn in enumerate(cols):
        pp = t.rows[0].cells[j].paragraphs[0]; pp.alignment = AL.CENTER
        _font(pp.add_run(str(cn)), "Times New Roman", 9, bold=True)
    for _, row in df.iterrows():
        cells = t.add_row().cells
        for j, cn in enumerate(cols):
            v = row[cn]
            if pd.isna(v): v = "—"
            elif isinstance(v, float): v = f"{v:.3f}".rstrip('0').rstrip('.') if abs(v) < 1000 else f"{v:.0f}"
            pp = cells[j].paragraphs[0]; pp.alignment = AL.CENTER if j > 0 else AL.LEFT
            _font(pp.add_run(str(v)), "Times New Roman", 9)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

# ============================== PORTADA ==============================
para("¿Cura la diabetes? Avances y desafíos del trasplante de islotes pancreáticos y la selección de "
     "candidatos a partir de datos reales de monitoreo continuo de glucosa",
     align=AL.CENTER, bold=True, size=14, space_after=4)
para("Does It Cure Diabetes? Advances and Challenges of Pancreatic Islet Transplantation and Candidate "
     "Selection Using Real-World Continuous Glucose Monitoring Data",
     align=AL.CENTER, italic=True, size=11, space_after=10)
para("Magali Bolivar, María Florencia Rossi, Matías Montiel, Nestor Balich, Franco Balich", align=AL.CENTER, space_after=2)
para("CAETI - Centro de Altos Estudios en Tecnología Informática", align=AL.CENTER, space_after=0)
para("Universidad Abierta Interamericana. Informática (UAI)", align=AL.CENTER, space_after=0)
para("Montes de Oca 745. Ciudad Autónoma de Buenos Aires, Argentina.", align=AL.CENTER, space_after=2)
para("{MagaliFlorencia.BolivarCruz, MariaFlorencia.Rossi, MatiasNicolas.MontielTorres}@alumnos.uai.edu.ar", align=AL.CENTER, italic=True, space_after=0)
para("{nestor.balich, francoadrian.balich}@uai.edu.ar", align=AL.CENTER, italic=True, space_after=10)

# ============================== RESUMEN ==============================
h1("Resumen")
para("El trasplante de islotes pancreáticos promete algo que hasta hace poco parecía inalcanzable: liberar "
"de la insulina a una persona con diabetes tipo 1. El Protocolo Edmonton y, más recientemente, inmunoterapias "
"como el tegoprubart han acercado esa promesa, pero a un costo —la inmunosupresión de por vida— que solo se "
"justifica en los casos más graves: quienes sufren hipoglucemias severas e inadvertidas y una glucemia "
"verdaderamente inestable. La pregunta práctica, entonces, no es solo si el trasplante cura, sino a quién "
"conviene ofrecérselo. Este trabajo la aborda con datos reales: se procesaron 14,8 millones de lecturas de "
"monitoreo continuo de glucosa (CGM) de 226 adultos con diabetes tipo 1 del estudio REPLACE-BG y, a partir de "
"la señal cruda, se reconstruyeron las métricas del consenso internacional —variabilidad, tiempo en rango y "
"tiempo en hipoglucemia— y los índices de riesgo validados de Kovatchev (LBGI, HBGI, ADRR), para traducir la "
"'labilidad' clínica en criterios objetivos y reproducibles. Aunque "
"el mal control resultó frecuente —dos de cada tres pacientes con variabilidad alta o tiempo en rango "
"insuficiente—, solo el 3,5 % reunió el perfil completo que justificaría evaluar un trasplante. Además, la "
"hipoglucemia inadvertida se asoció con los años de enfermedad y la edad, y no con las métricas de glucosa "
"del momento, lo que sugiere que la decisión no puede apoyarse únicamente en el CGM reciente. A ello se suma "
"un modelo de alerta temprana que, entrenado sobre unas 0,8 millones de ventanas de CGM y validado por "
"paciente, anticipa la hipoglucemia a 30 minutos con alta discriminación (AUC-ROC 0,98). El resultado es "
"una herramienta objetiva y reproducible para priorizar a los candidatos, que respalda con números la "
"necesidad de una selección estricta.", italic=True)
runs_para([("Palabras clave: ", True, False),
("trasplante de islotes; diabetes tipo 1; monitoreo continuo de glucosa; labilidad glucémica; "
"estratificación de candidatos; ciencia de datos.", False, True)])
h1("Abstract")
para("Pancreatic islet transplantation offers something long thought out of reach: freeing a person with "
"type 1 diabetes from insulin. The Edmonton Protocol and, more recently, immunotherapies such as tegoprubart "
"have brought that promise closer, but at a cost —lifelong immunosuppression— that is only justified in the "
"most severe cases: patients with severe, impaired-awareness hypoglycemia and truly unstable glucose. The "
"practical question is therefore not only whether transplantation cures, but for whom it is worth it. We "
"address it with real data: we processed 14.8 million continuous glucose monitoring (CGM) readings from 226 "
"adults with type 1 diabetes in the REPLACE-BG study and, from the raw signal, reconstructed the "
"international-consensus metrics —variability, time in range and time in hypoglycemia— and the validated "
"Kovatchev risk indices (LBGI, HBGI, ADRR), to turn clinical 'lability' into objective, reproducible "
"criteria. Although poor control was common —two of every three "
"patients had high variability or insufficient time in range— only 3.5% met the full profile that would "
"justify evaluating a transplant. Moreover, impaired awareness of hypoglycemia was associated with disease "
"duration and age rather than current glucose metrics, suggesting the decision cannot rest on recent CGM "
"alone. In addition, an early-warning model trained on ~0.8 million CGM windows and validated at the patient "
"level anticipates hypoglycemia 30 minutes ahead with high discrimination (AUC-ROC 0.98). The result is an "
"objective, reproducible tool to prioritize candidates that quantitatively supports strict selection.", italic=True)
runs_para([("Keywords: ", True, True),
("islet transplantation; type 1 diabetes; continuous glucose monitoring; glycemic lability; "
"candidate stratification; data science.", False, True)])

# ============================== 1. INTRODUCCIÓN ==============================
h1("1. Introducción")
para("La diabetes mellitus tipo 1 (DM1) se caracteriza por la destrucción autoinmune de las células beta "
"pancreáticas, lo que genera una dependencia absoluta de la insulina exógena. Pese a los avances en "
"insulinoterapia y en tecnología de administración, una subpoblación de pacientes experimenta labilidad "
"glucémica extrema, con hipoglucemias severas recurrentes e hipoglucemia inadvertida —un síndrome de falla "
"autonómica que eleva de forma drástica la morbimortalidad—. En este escenario, el trasplante alogénico de "
"islotes pancreáticos ha emergido como una intervención capaz de restaurar la masa de células beta y la "
"homeostasis glucémica [1], [2].", first_indent=0.5)
para("Sin embargo, el trasplante exige inmunosupresión crónica (habitualmente tacrolimus y sirolimus) con "
"riesgo de nefrotoxicidad, infecciones y neoplasias, por lo que se reserva para pacientes con inestabilidad "
"metabólica severa. Ensayos recientes con tegoprubart, un anticuerpo monoclonal anti-CD40L, muestran "
"resultados prometedores en independencia de insulina con menor toxicidad [3], [4], pero no eliminan la "
"premisa central: la relación riesgo-beneficio solo es favorable en "
"pacientes cuidadosamente seleccionados. Como concluye la literatura, en pacientes con buen control "
"metabólico y baja variabilidad glucémica el procedimiento presenta más riesgos que beneficios.", first_indent=0.5)
para("La selección de candidatos, no obstante, se apoya con frecuencia en criterios clínicos cualitativos. El "
"monitoreo continuo de glucosa (CGM) permite cuantificar de manera objetiva la variabilidad y la exposición "
"a hipoglucemia. El objetivo de este trabajo es doble y se apoya en una misma señal de CGM: por un lado, "
"construir con datos reales y un pipeline reproducible una estratificación objetiva de candidatos a "
"trasplante —cuantificando en una población de DM1 la prevalencia de los criterios de labilidad que definen "
"la indicación—; por otro, anticipar la hipoglucemia en el corto plazo para asistir la decisión en tiempo "
"real. Son dos horizontes de una misma pregunta: a quién conviene evaluar y cuándo conviene actuar.", first_indent=0.5)
para("El desafío no es menor del lado de los datos. Trabajar con casi 15 millones de registros de sensor "
"obliga a procesarlos en flujo, sin cargarlos enteros en memoria, y a reconstruir a partir de la señal cruda "
"las variables clínicas que realmente importan; recién sobre esa base tiene sentido buscar patrones y "
"modelar el riesgo. Ese recorrido —de la señal al criterio de candidatura— se documentó de principio a fin "
"para que cualquier equipo pueda repetirlo y auditarlo, y se concibió como apoyo a la decisión clínica: una "
"herramienta que ordena la evidencia, no que reemplaza el juicio del profesional.", first_indent=0.5)
para("En concreto, el trabajo aporta: (i) un procesamiento en flujo de 14,8 millones de lecturas que "
"reconstruye, desde la señal cruda, las métricas de consenso y los índices de riesgo validados de Kovatchev "
"(LBGI, HBGI, ADRR); (ii) una operacionalización transparente y auditable de los criterios de candidatura de "
"Edmonton, con análisis de sensibilidad de sus umbrales; y (iii) un modelo interpretable de alerta temprana "
"de hipoglucemia, validado por paciente para evitar la fuga de información. Los tres componentes comparten un "
"mismo principio: convertir una señal masiva en decisiones explicables que asisten —sin reemplazar— al "
"profesional.", first_indent=0.5)

# ============================== 2. MARCO CONCEPTUAL ==============================
h1("2. Marco conceptual")
para("El Protocolo Edmonton estableció la viabilidad del trasplante de islotes con un régimen "
"inmunosupresor libre de glucocorticoides [1] y fue validado en un ensayo internacional "
"multicéntrico [2]. Su indicación se reserva para la DM1 con hipoglucemias severas e "
"inadvertidas y labilidad metabólica [5], [6]. Los ensayos "
"de fase 3 del Clinical Islet Transplantation Consortium [7] y protocolos como el de la "
"Universidad de Illinois para la 'brittle diabetes' [8] consolidaron su eficacia sobre la "
"hipoglucemia, aunque la toxicidad inmunosupresora y la escasez de donantes siguen limitando su alcance "
"[9], [10].", first_indent=0.5)
para("La hipoglucemia inadvertida (impaired awareness of hypoglycemia, IAH) —la pérdida de los síntomas de "
"alarma— es el eje de la indicación: multiplica el riesgo de hipoglucemia severa. Por su parte, el consenso "
"internacional sobre interpretación del CGM [11] define objetivos estandarizados: tiempo "
"en rango (TIR 70–180 mg/dL) ≥ 70%, tiempo por debajo de rango (TBR <54 mg/dL) < 1% y coeficiente de "
"variación (%CV) < 36% como umbral de estabilidad glucémica. Estas métricas permiten traducir la 'labilidad' "
"clínica en variables objetivas y reproducibles, base de la estratificación que se propone.", first_indent=0.5)
para("El principal obstáculo del trasplante es la toxicidad de la inmunosupresión crónica. Para superarlo, "
"los ensayos de Eledon Pharmaceuticals (2024–2026) evaluaron el tegoprubart, un anticuerpo monoclonal "
"anti-CD40L que bloquea la coestimulación linfocitaria: en un estudio piloto con 12 adultos, la totalidad "
"alcanzó independencia de insulina con HbA1c inferior al 6,0% y buena tolerancia, sin nefrotoxicidad ni "
"infecciones oportunistas [3], [4]. Persisten, sin embargo, "
"limitaciones como la necesidad de infusiones intravenosas periódicas, la escasez de donantes y un "
"seguimiento a largo plazo aún en curso. En la Argentina, el procedimiento se enmarca en la normativa del "
"INCUCAI y se reserva para casos de DM1 con inestabilidad metabólica severa, lo que vuelve crítica una "
"estratificación objetiva de candidatos y el abordaje por equipos multidisciplinarios. En este contexto, "
"disponer de una herramienta cuantitativa y reproducible para priorizar la evaluación pretrasplante adquiere "
"particular relevancia.", first_indent=0.5)

# ============================== 3. METODOLOGÍA ==============================
h1("3. Metodología")
h2("3.1. Datos")
para("Se utilizó el conjunto de datos público REPLACE-BG [12] (Jaeb Center for Health "
"Research), un ensayo clínico aleatorizado en 226 adultos con DM1 de larga data y buen control basal "
"(HbA1c 6,4–9,0%). El dataset incluye el registro completo de CGM (14,8 millones de lecturas), datos de "
"cribado (demografía, antecedentes), determinaciones de HbA1c y cuestionarios de percepción de hipoglucemia. "
"El acceso es de descarga directa para investigación.")
h2("3.2. Métricas de monitoreo continuo de glucosa")
para("A partir del archivo crudo de CGM se calcularon, por paciente, las métricas del consenso internacional "
"[11]: glucosa media, coeficiente de variación (%CV = desvío estándar / media), "
"indicador de gestión de la glucosa (GMI), tiempo en rango (TIR 70–180 mg/dL), tiempo por debajo de rango "
"(TBR <70 y <54 mg/dL) y tiempo por encima de rango (TAR >180 y >250 mg/dL). El procesamiento se realizó en "
"flujo (streaming) sobre el archivo de 837 MB, acumulando estadísticos sin cargarlo íntegramente en memoria, "
"lo que garantiza la reproducibilidad y la escalabilidad del pipeline (Python, pandas).")
para("Además, a partir de la misma señal cruda se calcularon los índices de riesgo glucémico de Kovatchev "
"[15]: el Low Blood Glucose Index (LBGI) y el High Blood Glucose Index (HBGI) —que cuantifican por separado "
"el riesgo de hipoglucemia y de hiperglucemia mediante una transformación simétrica de la escala de glucosa— "
"y el Average Daily Risk Range (ADRR), que resume la labilidad diaria combinando los extremos bajo y alto de "
"cada día. Estos índices son la versión apropiada para CGM y validada del Lability Index y el HYPO score que "
"el grupo de Edmonton (Ryan et al. [16]) empleó con glucemias capilares para evaluar la severidad de la "
"hipoglucemia y la labilidad en la selección de candidatos a trasplante.")
h2("3.3. Criterios de candidatura y análisis")
para("Los criterios de la indicación de trasplante se operacionalizaron combinando los cuestionarios y las "
"métricas de CGM: (i) hipoglucemia inadvertida (IAH), definida por la pérdida de los síntomas de alarma; "
"(ii) variabilidad glucémica alta (%CV ≥ 36%); (iii) exposición excesiva a hipoglucemia (TBR <54 > 1%); y "
"(iv) tiempo en rango insuficiente (TIR < 70%). El perfil de labilidad glucémica —proxy de la candidatura de "
"Edmonton— se definió como la coexistencia de hipoglucemia inadvertida y exposición excesiva a hipoglucemia. "
"El análisis incluyó estadística descriptiva, comparación entre subgrupos (prueba t de Welch), segmentación "
"no supervisada de fenotipos glucémicos mediante estandarización, PCA y K-Means (k=3) con scikit-learn [13], "
"y una regresión logística de la hipoglucemia inadvertida con statsmodels [14], en función de las métricas de "
"CGM, la duración de la enfermedad y la edad. Por último, se evaluó la robustez del criterio de candidatura "
"mediante un análisis de sensibilidad que recalcula su prevalencia al variar el umbral operacional de "
"exposición a hipoglucemia.")
h2("3.4. Alerta temprana de hipoglucemia")
para("Como componente de apoyo a la decisión se planteó, además, un modelo de anticipación de la "
"hipoglucemia —un problema reconocido de la analítica de CGM (Oviedo et al. [17])—. En lugar de tratar cada "
"paciente como una unidad, se reencuadró la señal en ventanas deslizantes: a partir de los 60 minutos previos "
"de CGM se predice si la glucosa caerá por debajo de 70 mg/dL en los siguientes 30 minutos. De cada ventana "
"se extrajeron variables interpretables (glucosa actual, media, desvío, mínimo, máximo, pendiente, variación, "
"proporción de lecturas bajas y hora del día), descartando las que cruzaban huecos del sensor, lo que "
"produjo cerca de 0,8 millones de ejemplos etiquetados. Se compararon una línea base ingenua (solo la glucosa "
"actual), una regresión logística y un modelo de gradient boosting (scikit-learn [13]), evaluados con "
"partición POR PACIENTE —ningún paciente comparte ventanas entre entrenamiento y test, para evitar la fuga de "
"información— y validación cruzada por grupos. Se reportan el área bajo la curva ROC y la de precisión-recall, "
"más apropiada ante el desbalance de clases, y la importancia de permutación como lente de interpretabilidad.")

# ============================== 4. RESULTADOS ==============================
h1("4. Resultados")
h2("4.1. Características de la cohorte")
para("La cohorte (N=226) tenía una edad media de 44,0 años, una duración media de la DM1 de 23,3 años y una "
"HbA1c basal media de 7,3%, con equilibrio por sexo (Tabla 1). Pese al buen control aparente, las métricas "
"de CGM revelaron una carga glucémica considerable: TIR medio del 63,1%, %CV medio del 37,4% y un tiempo en "
"hipoglucemia <54 mg/dL del 1,0% de media. Los índices de riesgo (Tabla 1) situaron a la cohorte en un LBGI "
"medio de 1,0 —riesgo de hipoglucemia bajo, coherente con una población que excluyó la hipoglucemia severa— "
"y un HBGI medio de 7,6, dominado por la hiperglucemia.")
tabla_csv(TAB/"tabla1_cohorte.csv", "Tabla 1. Características clínicas y métricas de CGM de la cohorte (N=226).")
h2("4.2. Prevalencia de los criterios de candidatura")
para("La mayoría de la cohorte no alcanzaba los objetivos de consenso: el 71,7% presentaba un TIR "
"insuficiente y el 64,6% una variabilidad glucémica alta (Figura 1, Tabla 2). La exposición excesiva a "
"hipoglucemia afectaba al 34,5% y la hipoglucemia inadvertida al 17,3%. Sin embargo, al exigir el perfil "
"completo de labilidad (hipoglucemia inadvertida más exposición excesiva a hipoglucemia), solo el 3,5% de la "
"cohorte (8 pacientes) reunía la condición compatible con la indicación de trasplante (intervalo de "
"confianza del 95% de Wilson: 1,8–6,8%). La Figura 2 muestra "
"las distribuciones de las métricas clave frente a los umbrales de consenso.")
tabla_csv(TAB/"tabla2_prevalencia_criterios.csv", "Tabla 2. Prevalencia de los criterios de candidatura en la cohorte.")
figura(FIG/"fig1_prevalencia_criterios.png", "Figura 1. Prevalencia de los criterios de candidatura a trasplante de islotes; en naranja, el perfil de labilidad completo.", width=15)
figura(FIG/"fig2_distribuciones.png", "Figura 2. Distribución de las métricas de CGM (%CV, TBR<54 y TIR) frente a los objetivos del consenso internacional [11].", width=16)
h2("4.3. Mapa de riesgo y fenotipos glucémicos")
para("El mapa de riesgo (Figura 3) evidencia una relación positiva entre la variabilidad glucémica y la "
"exposición a hipoglucemia: los pacientes con perfil de labilidad se concentran en el cuadrante de %CV y "
"TBR<54 elevados. La segmentación no supervisada (Figura 4, Tabla 4) identificó tres fenotipos glucémicos: "
"uno de buen control (mayor TIR, menor variabilidad), uno hiperglucémico (glucosa media y TAR elevados con "
"TIR bajo) y uno lábil (alta variabilidad y mayor exposición a hipoglucemia), que concentra los perfiles de "
"candidatura. La comparación entre el subgrupo con perfil de labilidad y el resto (Tabla 3) confirma "
"diferencias marcadas en variabilidad y tiempo en hipoglucemia; en particular, el LBGI —índice validado de "
"riesgo de hipoglucemia— fue significativamente mayor en el grupo de labilidad (1,4 frente a 1,0; p=0,008), "
"lo que ancla el criterio categórico en una medida continua y validada del riesgo.")
figura(FIG/"fig3_mapa_riesgo.png", "Figura 3. Mapa de riesgo glucémico (variabilidad %CV vs. tiempo en hipoglucemia <54 mg/dL); líneas discontinuas: umbrales de consenso.", width=12)
tabla_csv(TAB/"tabla3_labilidad_vs_resto.csv", "Tabla 3. Comparación entre el subgrupo con perfil de labilidad y el resto de la cohorte (prueba t de Welch).")
figura(FIG/"fig4_fenotipos_pca.png", "Figura 4. Fenotipos glucémicos identificados por K-Means (k=3) en el espacio de las dos primeras componentes principales.", width=11)
tabla_csv(TAB/"tabla4_fenotipos.csv", "Tabla 4. Perfil promedio de los fenotipos glucémicos (K-Means, k=3).")
figura(FIG/"fig5_comparacion_boxplots.png", "Figura 5. Comparación de las métricas de CGM entre el subgrupo con perfil de labilidad y el resto de la cohorte (líneas discontinuas: umbrales de consenso).", width=15)
para("La matriz de correlación entre las métricas de CGM (Figura 6) confirma la coherencia interna de los "
"indicadores: la variabilidad (%CV) se asocia positivamente con la exposición a hipoglucemia (TBR<54) y "
"negativamente con el tiempo en rango, mientras que el score de riesgo glucémico resume de forma compacta "
"esta estructura, correlacionando de manera fuerte con las variables que penalizan el control.")
figura(FIG/"fig6_correlacion.png", "Figura 6. Matriz de correlación de Pearson entre las métricas de CGM y el score de riesgo glucémico.", width=12)
h2("4.4. Determinantes de la hipoglucemia inadvertida")
para("La regresión logística (Tabla 5) mostró que la hipoglucemia inadvertida se asocia de manera "
"significativa con la duración de la DM1 (OR≈1,04 por año; p=0,016) y con la edad (OR≈1,04; p=0,039), pero "
"no con las métricas de CGM del período de observación (%CV, TBR<54, TIR). Este resultado es coherente con la "
"naturaleza crónica y autonómica de la falla en la percepción de hipoglucemia: la IAH refleja años de "
"exposición acumulada más que el estado glucémico puntual, lo que refuerza la necesidad de combinar "
"cuestionarios clínicos con métricas objetivas para identificar candidatos. El modelo es globalmente "
"significativo (test de razón de verosimilitud, p<0,001) y no presenta multicolinealidad (VIF < 2,4 en todos "
"los predictores); la significación de la duración y la edad se mantiene bajo errores estándar robustos "
"(HC1), lo que respalda la solidez del resultado. Con 39 eventos de hipoglucemia inadvertida para cinco "
"predictores —cerca de ocho eventos por variable—, el modelo se mantiene dentro de los márgenes habituales "
"de estabilidad, aunque sus estimaciones deben leerse con la prudencia propia de una muestra acotada.")
tabla_csv(TAB/"tabla5_logit_iah.csv", "Tabla 5. Regresión logística de la hipoglucemia inadvertida (IAH): odds ratios, p-valores clásicos y robustos (HC1) y VIF.")

h2("4.5. Índices de riesgo validados y robustez de los criterios")
para("Para anclar la definición operacional en medidas validadas, se calcularon los índices de riesgo de "
"Kovatchev sobre la señal cruda (Figura 7). El grupo con perfil de labilidad mostró un LBGI —riesgo de "
"hipoglucemia— significativamente mayor (p=0,008) y un ADRR más alto, mientras que su HBGI fue menor, un "
"patrón coherente con un riesgo desplazado hacia la hipoglucemia. Cabe señalar que, calculado sobre CGM de "
"alta frecuencia, el ADRR alcanza valores superiores a los de los cortes originales de Kovatchev —derivados "
"de cuatro glucemias capilares por día—, por lo que estos índices se interpretan aquí de forma comparativa "
"entre grupos y no contra umbrales absolutos.")
para("La prevalencia de la candidatura depende, como toda operacionalización, de los umbrales elegidos. El "
"análisis de sensibilidad (Tabla 6) muestra que, al mover el umbral de exposición a hipoglucemia del 0,5% al "
"2% de TBR<54, la prevalencia del perfil pasa del 8,8% al 0,4%, correspondiendo el 3,5% al umbral de consenso "
"(1%); exigir además variabilidad alta (%CV≥36%) casi no altera el resultado. Hacer explícitos estos "
"umbrales —y poder recalcular el criterio— es parte del aporte: convierte una decisión tradicionalmente "
"cualitativa en un procedimiento transparente y auditable.")
figura(FIG/"fig7_indices_riesgo.png", "Figura 7. Índices de riesgo glucémico de Kovatchev (LBGI, HBGI, ADRR) según el perfil de labilidad.", width=15)
tabla_csv(TAB/"tabla7_sensibilidad.csv", "Tabla 6. Sensibilidad de la prevalencia de candidatura al umbral operacional de hipoglucemia (TBR<54).")

h2("4.6. Alerta temprana de hipoglucemia")
para("El reencuadre en ventanas generó 795.851 ejemplos (prevalencia de hipoglucemia a 30 minutos del 5,8%). "
"Sobre pacientes no vistos en el entrenamiento, el modelo de gradient boosting alcanzó un AUC-ROC de 0,985 y "
"un AUC-PR de 0,857, y se mantuvo estable en validación cruzada por grupos (AUC-ROC 0,984 ± 0,001), lo que "
"confirma que no hay fuga de información (Tabla 7, Figura 8). Es honesto señalar que la línea base ingenua —la "
"glucosa actual— ya es un predictor fuerte (AUC-ROC 0,974), porque el nivel presente condiciona el de los "
"próximos minutos; el aporte del aprendizaje automático se ve sobre todo en la precisión-recall (AUC-PR del "
"0,803 al 0,857), la métrica que gobierna las falsas alarmas en un sistema de alerta. La importancia de "
"permutación (Figura 9) muestra que el modelo se apoya en la glucosa actual, la media reciente, la tendencia "
"y la variabilidad de la ventana —variables clínicamente interpretables—, lo que lo vuelve utilizable como "
"apoyo transparente y no como una caja negra.")
figura(FIG/"fig8_roc_pr.png", "Figura 8. Alerta temprana de hipoglucemia a 30 minutos: curvas ROC y precisión-recall en pacientes no vistos.", width=16)
tabla_csv(TAB/"tabla8_early_warning.csv", "Tabla 7. Desempeño de los modelos de alerta temprana de hipoglucemia (test por paciente).")
figura(FIG/"fig9_importancia_ew.png", "Figura 9. Importancia de permutación: variables que más pesan en la alerta de hipoglucemia.", width=12)

# ============================== 5. DISCUSIÓN ==============================
h1("5. Discusión")
para("Los resultados cuantifican una intuición clínica central del trasplante de islotes: aunque la carga "
"glucémica subóptima es muy frecuente en la DM1 (dos de cada tres pacientes con variabilidad alta o TIR "
"insuficiente), el perfil de labilidad severa que justifica el riesgo de la inmunosupresión es minoritario "
"(3,5%). Esto respalda de forma empírica la recomendación de reservar el procedimiento —y las nuevas terapias "
"como el tegoprubart— para un subgrupo cuidadosamente seleccionado, y advierte contra una indicación "
"guiada solo por el mal control global.", first_indent=0.5)
para("Ese 3,5% ayuda además a dimensionar el problema. Proyectado sobre el conjunto de personas con diabetes "
"tipo 1, incluso una fracción pequeña representa un número nada despreciable de candidatos potenciales que "
"hoy podrían no estar siendo identificados de manera sistemática. Al mismo tiempo, la estrechez del grupo "
"explica por qué los programas de trasplante operan con listas reducidas y por qué una selección laxa, "
"guiada solo por el mal control, expondría a muchos pacientes a la toxicidad de la inmunosupresión sin un "
"beneficio proporcional. El valor del enfoque propuesto es, precisamente, ofrecer un criterio cuantitativo y "
"repetible para trazar esa línea con transparencia, allí donde la práctica todavía se apoya en gran medida "
"en la valoración cualitativa.", first_indent=0.5)
para("El hallazgo de que la hipoglucemia inadvertida depende de la duración y la edad, y no de las métricas "
"de CGM del momento, tiene una implicancia práctica: la estratificación de candidatos no puede basarse "
"únicamente en el CGM reciente, sino que debe integrar la historia clínica y la evaluación de la percepción "
"de hipoglucemia. La combinación de ambas fuentes —como propone este pipeline— ofrece una herramienta "
"objetiva y reproducible para priorizar la evaluación pretrasplante.", first_indent=0.5)
para("El momento en que se plantea esta pregunta no es casual. Las nuevas inmunoterapias —el tegoprubart y "
"otros esquemas libres de inhibidores de calcineurina— están corriendo la frontera de lo posible: prometen "
"la independencia de insulina con menor toxicidad y, si esa promesa se confirma, ampliarán el universo de "
"pacientes para quienes el trasplante sea una opción razonable. Paradójicamente, eso vuelve más necesaria —y "
"no menos— una selección rigurosa: cuanto más se acerque el trasplante a una verdadera cura, más importará "
"decidir con transparencia quién se beneficia y en qué momento. Un criterio reproducible como el que aquí se "
"propone busca aportar a esa discusión, que todavía descansa en gran medida en la experiencia clínica "
"individual.", first_indent=0.5)
para("Entre las limitaciones, REPLACE-BG es una población de DM1 con buen control basal en la que la"
"hipoglucemia severa fue criterio de exclusión, por lo que la prevalencia del perfil de labilidad "
"probablemente subestima la de una población con indicación real de trasplante; el diseño es transversal y "
"la definición operacional de candidatura es un proxy de los criterios clínicos completos. No obstante, el "
"marco metodológico es directamente transferible a cohortes de mayor riesgo.", first_indent=0.5)

# ============================== 6. CONCLUSIONES ==============================
h1("6. Conclusiones y trabajos futuros")
para("Se presentó un pipeline reproducible de ciencia de datos que, sobre 14,8 millones de lecturas reales "
"de CGM, estratifica objetivamente a pacientes con DM1 según los criterios de candidatura a trasplante de "
"islotes. La variabilidad glucémica alta y el tiempo en rango insuficiente son muy frecuentes, pero el "
"perfil de labilidad severa es minoritario (3,5%), lo que respalda una selección estricta de candidatos. La "
"hipoglucemia inadvertida se asoció a la duración de la enfermedad y a la edad más que al CGM del momento. "
"Más allá del hallazgo clínico, el trabajo deja un procedimiento reproducible —de la señal cruda a los "
"índices de riesgo validados y al criterio de candidatura— que puede auditarse y reutilizarse en otras "
"cohortes, sumado a un modelo interpretable de alerta temprana que anticipa la hipoglucemia a 30 minutos "
"(validado por paciente). Todo ello se plantea como apoyo a la "
"decisión, en sintonía con el factor humano de la inteligencia artificial (Humanware 5.0): una herramienta "
"que asiste, sin reemplazar, el juicio del profesional.", first_indent=0.5)
para("Como líneas futuras se plantea: (i) aplicar el pipeline a cohortes con indicación real de trasplante "
"(p. ej., registros CIT/CITR mediante solicitud de acceso); (ii) extender el modelo de alerta a la "
"hipoglucemia severa (<54 mg/dL) y a horizontes más largos, e incorporar índices adicionales como el MAGE; "
"(iii) integrar datos de los "
"nuevos ensayos con tegoprubart para vincular el perfil glucémico basal con la respuesta post-trasplante; y "
"(iv) desarrollar una herramienta de tablero interactivo para la evaluación pretrasplante.", first_indent=0.5)

# ============================== REFERENCIAS ==============================
h1("Referencias")
refs = [
'[1] A. M. J. Shapiro et al., "Islet transplantation in seven patients with type 1 diabetes mellitus using a glucocorticoid-free immunosuppressive regimen", New England Journal of Medicine, vol. 343, n.º 4, pp. 230–238, 2000.',
'[2] A. M. J. Shapiro, C. Ricordi, B. J. Hering et al., "International Trial of the Edmonton Protocol for Islet Transplantation", New England Journal of Medicine, vol. 355, n.º 13, pp. 1318–1330, 2006.',
'[3] N. M. Paucara Saavedra, "¿Cura a la diabetes? Avances y desafíos en el trasplante de islotes pancreáticos", XVI Congreso Argentino de Estudiantes de Nutrición, Facultad de Ciencias Médicas, Universidad de Buenos Aires, 2026.',
'[4] M. Cassola y O. J. Leal Niebla, "Tegoprubart and the CD40L Pathway: Promise and Remaining Questions in CNI-Free Transplantation", BioNatura Journal, 2026.',
'[5] Health Quality Ontario, "Pancreas Islet Transplantation for Patients With Type 1 Diabetes Mellitus: A Clinical Evidence Review", Ontario Health Technology Assessment Series, vol. 15, n.º 16, pp. 1–84, 2015.',
'[6] M. R. Rickels y R. P. Robertson, "Pancreatic Islet Transplantation in Humans: Recent Progress and Future Directions", Endocrine Reviews, vol. 40, n.º 2, pp. 631–668, 2019.',
'[7] C. Ricordi, J. S. Goldstein, A. N. Balamurugan et al., "NIH-sponsored Clinical Islet Transplantation Consortium Phase 3 Trial", Diabetes, vol. 65, pp. 3418–3428, 2016.',
'[8] A. Gangemi et al., "Islet Transplantation for Brittle Type 1 Diabetes: The UIC Protocol", American Journal of Transplantation, vol. 8, pp. 1250–1261, 2008.',
'[9] A. Bruni et al., "Islet cell transplantation for the treatment of type 1 diabetes: recent advances and future challenges", Diabetes, Metabolic Syndrome and Obesity, vol. 7, pp. 211–223, 2014.',
'[10] A. M. J. Shapiro, "Islet Transplantation in Type 1 Diabetes: Ongoing Challenges, Refined Procedures, and Long-Term Outcome", The Review of Diabetic Studies, vol. 9, n.º 4, pp. 385–406, 2012.',
'[11] T. Battelino, T. Danne, R. M. Bergenstal et al., "Clinical Targets for Continuous Glucose Monitoring Data Interpretation: Recommendations From the International Consensus on Time in Range", Diabetes Care, vol. 42, n.º 8, pp. 1593–1603, 2019.',
'[12] R. W. Beck, T. D. Riddlesworth, K. Ruedy et al., "Continuous Glucose Monitoring Versus Usual Care in Patients With Type 1 Diabetes (REPLACE-BG)", Diabetes Care, vol. 40, n.º 4, pp. 538–545, 2017.',
'[13] F. Pedregosa et al., "Scikit-learn: Machine Learning in Python", Journal of Machine Learning Research, vol. 12, pp. 2825–2830, 2011.',
'[14] S. Seabold y J. Perktold, "Statsmodels: Econometric and statistical modeling with Python", en Proc. 9th Python in Science Conf., 2010.',
'[15] B. P. Kovatchev, E. Otto, D. Cox, L. Gonder-Frederick y W. Clarke, "Evaluation of a New Measure of Blood Glucose Variability in Diabetes", Diabetes Care, vol. 29, n.º 11, pp. 2433–2438, 2006.',
'[16] E. A. Ryan, T. Shandro, K. Green et al., "Assessment of the Severity of Hypoglycemia and Glycemic Lability in Type 1 Diabetic Subjects Undergoing Islet Transplantation", Diabetes, vol. 53, n.º 4, pp. 955–962, 2004.',
'[17] S. Oviedo, J. Vehí, R. Calm y J. Armengol, "A review of personalized blood glucose prediction strategies for T1DM patients", International Journal for Numerical Methods in Biomedical Engineering, vol. 33, n.º 6, e2833, 2017.',
]
for r in refs:
    p = doc.add_paragraph(); p.alignment = AL.JUSTIFY; p.paragraph_format.space_after = Pt(3)
    _font(p.add_run(r), "Times New Roman", 10)

doc.save(OUT)
print("Guardado:", OUT)
